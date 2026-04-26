"""
Works routes — Work CRUD, ingest, modules (serializer segments), bulk import, asset generation.
"""
import io
import os
import re
import json
import uuid
import threading
from datetime import datetime, timedelta, timezone

from flask import Blueprint, jsonify, request
from utils.json_store import read_json, write_json
from utils.constants import (
    PROMO_WORKS_FILE, PROMO_MESSAGES_FILE, PROMO_ASSETS_FILE,
    PROMO_SETTINGS_FILE, DEFAULT_PROMO_SETTINGS,
    BOOK_SERIALIZER_SYSTEM_PROMPT, BOOK_SERIALIZER_FIXED_PROMPT,
    CHAPTER_SYNOPSIS_PROMPT, CHAPTER_TAGLINE_PROMPT,
    CHAPTER_BLURB_PROMPT, HEADER_IMAGE_PROMPT_SYSTEM,
    CONTENT_PIPELINE_FILE,
)
from services.ai_service import call_ai
from services.asset_manager import create_asset, update_asset, get_asset

bp = Blueprint('works', __name__)

# ── Module-level state ────────────────────────────────────────────────────────

_PENDING_SPLITS: dict = {}   # work_id → list[chapter_dict]  (ephemeral, per-session)
BULK_JOBS: dict        = {}   # job_id  → status dict


# ── Helpers ───────────────────────────────────────────────────────────────────

def _detect_chapter_line(line: str):
    """Return the line stripped if it looks like a chapter heading, else None."""
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return None
    # "Chapter 1", "Chapter IV", "CHAPTER One", optionally followed by ": Subtitle"
    if re.match(
        r'^(?:CHAPTER|Chapter|chapter)\s+'
        r'(?:\d+|[IVXLCDM]+|'
        r'One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|'
        r'Eleven|Twelve|Thirteen|Fourteen|Fifteen|'
        r'Sixteen|Seventeen|Eighteen|Nineteen|Twenty|'
        r'Thirty|Forty|Fifty|Sixty|Seventy|Eighty|Ninety|'
        r'(?:Twenty|Thirty|Forty|Fifty|Sixty|Seventy|Eighty|Ninety)'
        r'(?:-| )'
        r'(?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine))'
        r'(?:\s*[:\-\u2013\u2014]\s*.+)?$',
        stripped
    ):
        return stripped
    # "1." or "1. Title" — digit-dot, at most 8 words
    if re.match(r'^\d{1,3}\.\s*(?:\S.*)?$', stripped) and len(stripped.split()) <= 8:
        return stripped
    # Bare number on its own line
    if re.match(r'^\d{1,3}$', stripped):
        return stripped
    return None


_ASSET_PROMPTS = {
    'synopsis': CHAPTER_SYNOPSIS_PROMPT,
    'tagline':  CHAPTER_TAGLINE_PROMPT,
    'blurb':    CHAPTER_BLURB_PROMPT,
}


def _run_bulk_generate(job_id: str, work_id: str,
                       module_items: list, asset_types: list):
    """
    Background thread: generate assets sequentially, update BULK_JOBS[job_id].
    module_items: list of {id, title, content} — source-agnostic.
    """
    job = BULK_JOBS[job_id]
    total = len(module_items) * len(asset_types)
    job['total'] = total
    done = 0

    for item in module_items:
        module_id = item['id']
        title     = item.get('title') or f"Module {module_id[:8]}"
        content   = item.get('content', '')

        for asset_type in asset_types:
            if job.get('cancelled'):
                job['status'] = 'cancelled'
                return

            asset_id = f"asset_{asset_type}_{module_id}"
            existing = get_asset(asset_id)
            if existing and existing.get('content', '').strip():
                done += 1
                job['done']    = done
                job['current'] = f"Skipped {title[:40]} / {asset_type} (already exists)"
                continue

            job['current'] = f"Generating {asset_type} for {title[:40]}..."
            system_prompt  = _ASSET_PROMPTS.get(asset_type, '')
            user_msg       = f"Chapter Title: {title}\n\nChapter Content:\n{content[:3000]}"

            try:
                result = call_ai("work_serializer", [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_msg},
                ], max_tokens=400)

                create_asset({
                    "id":         asset_id,
                    "type":       asset_type,
                    "title":      f"{title} \u2014 {asset_type.title()}",
                    "work_id":    work_id,
                    "module_id":  module_id,
                    "content":    result.strip(),
                    "production": "done",
                })
            except Exception as e:
                job['errors'].append(f"{title[:30]}/{asset_type}: {str(e)}")

            done += 1
            job['done'] = done

    job['status']       = 'done'
    job['completed_at'] = datetime.utcnow().isoformat() + "Z"


def _sync_chunk_statuses(works_data):
    """
    Cross-reference chunk message IDs against promo_messages.json and update
    chunk statuses in-place. Returns True if any status changed (caller should
    write back to disk).
    """
    msgs  = (read_json(PROMO_MESSAGES_FILE) or {}).get('messages', [])
    lookup = {m['id']: m.get('status', 'queued') for m in msgs}
    changed = False

    for work in works_data.get('works', []):
        for chunk in work.get('chunks', []):
            old_status = chunk.get('status', 'pending')
            if old_status in ('pending',):
                continue  # nothing to sync

            # Novel dual-delivery: use the channel message as the authoritative status
            channel_id = chunk.get('channel_message_id')
            vip_id     = chunk.get('vip_message_id')
            single_id  = chunk.get('message_id')

            if channel_id and channel_id in lookup:
                new_status = lookup[channel_id]
            elif vip_id and vip_id in lookup:
                # Only VIP tracked: use its status
                new_status = lookup[vip_id]
            elif single_id and single_id in lookup:
                new_status = lookup[single_id]
            else:
                continue  # message not in outbox — leave as-is

            if new_status != old_status:
                chunk['status'] = new_status
                changed = True

    return changed


@bp.route('/api/works', methods=['GET'])
def list_works():
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    if _sync_chunk_statuses(works_data):
        write_json(PROMO_WORKS_FILE, works_data)
    return jsonify(works_data)


@bp.route('/api/works', methods=['POST'])
def create_work():
    data  = request.get_json()
    title = data.get('title', '').strip()
    if not title:
        return jsonify({"error": "title is required"}), 400

    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    now        = datetime.utcnow().isoformat() + "Z"
    pipeline_module_id = data.get('pipeline_module_id', '').strip()
    website_url        = data.get('website_url', '').strip()

    # If no website_url provided but a pipeline_module_id was given, derive it
    if not website_url and pipeline_module_id:
        pipeline = read_json(CONTENT_PIPELINE_FILE) or []
        pm = next((e for e in pipeline if e['id'] == pipeline_module_id), None)
        if pm:
            pub_info    = pm.get('website_publish_info') or {}
            website_url = pub_info.get('chapter_url', '')

    new_work   = {
        "id":          str(uuid.uuid4()),
        "title":       title,
        "author":      data.get('author', ''),
        "patreon_url": data.get('patreon_url', ''),
        "website_url": website_url,
        "chunks":      [],
        "created_at":  now,
        "updated_at":  now,
    }
    if pipeline_module_id:
        new_work["pipeline_module_id"] = pipeline_module_id
    works_data["works"].append(new_work)
    write_json(PROMO_WORKS_FILE, works_data)
    return jsonify(new_work), 201


@bp.route('/api/works/<work_id>', methods=['GET'])
def get_work(work_id):
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404
    return jsonify(work)


@bp.route('/api/works/<work_id>', methods=['PUT'])
def update_work(work_id):
    data       = request.get_json()
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work_idx   = next((i for i, w in enumerate(works_data.get("works", [])) if w['id'] == work_id), None)
    if work_idx is None:
        return jsonify({"error": "Work not found"}), 404
    work = works_data["works"][work_idx]
    for field in ['title', 'author', 'patreon_url', 'website_url', 'pipeline_module_id', 'post_header']:
        if field in data:
            work[field] = data[field]
    work['updated_at'] = datetime.utcnow().isoformat() + "Z"
    write_json(PROMO_WORKS_FILE, works_data)
    return jsonify(work)


@bp.route('/api/works/<work_id>', methods=['DELETE'])
def delete_work_route(work_id):
    from services.asset_manager import delete_work
    delete_work(work_id)
    return jsonify({"success": True})


@bp.route('/api/works/<work_id>/ingest', methods=['POST'])
def ingest_work_content(work_id):
    text                  = ""
    target_words_override = None
    max_words_override    = None
    body                  = {}

    if request.content_type and 'multipart' in request.content_type:
        uploaded              = request.files.get('file')
        target_words_override = request.form.get('target_words')
        max_words_override    = request.form.get('max_words')

        if not uploaded:
            return jsonify({"error": "No file provided."}), 400

        filename = uploaded.filename.lower()
        if filename.endswith('.txt'):
            text = uploaded.read().decode('utf-8', errors='replace')
        elif filename.endswith('.rtf'):
            raw  = uploaded.read().decode('utf-8', errors='replace')
            text = re.sub(r'\\(par|line)\b\s*', '\n', raw)
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'[{}\\]', '', text)
            text = re.sub(r'[ \t]+', ' ', text).strip()
        elif filename.endswith('.docx'):
            try:
                from docx import Document as DocxDocument
                uploaded.seek(0)
                doc  = DocxDocument(io.BytesIO(uploaded.read()))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as e:
                return jsonify({"error": f"Failed to read DOCX: {e}"}), 400
        else:
            return jsonify({"error": "Unsupported file type. Use .txt, .rtf, or .docx"}), 400
    else:
        body                  = request.get_json() or {}
        text                  = body.get('text', '').strip()
        target_words_override = body.get('target_words')
        max_words_override    = body.get('max_words')  # kept for fallback

    if not text.strip():
        return jsonify({"error": "Input text is empty."}), 400

    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    # Resolve profile: prefer profile_id on the work, then fall back to request body
    settings = read_json(PROMO_SETTINGS_FILE) or DEFAULT_PROMO_SETTINGS
    profiles = settings.get('serializer_profiles') or []
    profile  = None

    profile_id = (body.get('profile_id') if not (request.content_type and 'multipart' in request.content_type) else None) \
                 or work.get('profile_id')
    if profile_id:
        profile = next((p for p in profiles if p['id'] == profile_id), None)

    num_chunks   = profile['num_chunks']   if profile else None
    target_words = int(profile['target_words'] if profile else (target_words_override or 300))

    now                = datetime.utcnow().isoformat() + "Z"
    new_chunks         = []
    errors             = []
    current_part_count = len(work.get("chunks", [])) + 1

    if num_chunks:
        # ── Fixed-chunk mode: one AI call, split into exactly N chunks ──────────
        system_prompt = BOOK_SERIALIZER_FIXED_PROMPT.format(
            work_title=work.get('title', 'Untitled Work'),
            start_part=current_part_count,
            num_chunks=num_chunks,
            target_words=target_words,
        )
        response_text = ""
        try:
            response_text = call_ai(
                "work_serializer",
                [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": f"Story text to split:\n\n{text}"},
                ],
                max_tokens=4000,
                response_format={"type": "json_object"}
            )
            clean = response_text.strip()
            clean = re.sub(r'^```(?:json)?\s*', '', clean, flags=re.MULTILINE)
            clean = re.sub(r'\s*```\s*$',       '', clean, flags=re.MULTILINE)
            data_json    = json.loads(clean.strip())
            segment_list = data_json.get("chunks", [])
            for c in segment_list:
                content = c.get("content", "").replace('\\n', '\n')
                new_chunks.append({
                    "id":               str(uuid.uuid4()),
                    "content":          content,
                    "cliffhanger_note": c.get("cliffhanger_note", ""),
                    "status":           "pending",
                    "word_count":       len(content.split()),
                    "created_at":       now,
                })
        except Exception as e:
            errors.append(str(e))

    else:
        # ── Auto mode: sliding windows, AI decides chunk count per window ────────
        window_size    = 600
        lines          = text.split('\n')
        windows        = []
        current_lines  = []
        current_wcount = 0

        for line in lines:
            line_wcount = len(line.split())
            if current_wcount + line_wcount > window_size and current_lines:
                windows.append('\n'.join(current_lines))
                current_lines  = [line]
                current_wcount = line_wcount
            else:
                current_lines.append(line)
                current_wcount += line_wcount
        if current_lines:
            windows.append('\n'.join(current_lines))

        if not windows:
            return jsonify({"error": "Input text produced no processable windows."}), 400

        for idx, window in enumerate(windows):
            system_prompt = BOOK_SERIALIZER_SYSTEM_PROMPT.format(
                work_title=work.get('title', 'Untitled Work'),
                start_part=current_part_count,
                max_words=target_words,
            )
            response_text = ""
            try:
                response_text = call_ai(
                    "work_serializer",
                    [
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": f"Text to serialize:\n\n{window}"},
                    ],
                    max_tokens=4000,
                    response_format={"type": "json_object"}
                )
                clean = response_text.strip()
                clean = re.sub(r'^```(?:json)?\s*', '', clean, flags=re.MULTILINE)
                clean = re.sub(r'\s*```\s*$',       '', clean, flags=re.MULTILINE)
                clean = clean.strip()
                try:
                    data_json = json.loads(clean)
                except json.JSONDecodeError:
                    repaired = re.sub(
                        r'("(?:[^"\\]|\\.)*")',
                        lambda m: m.group(0).replace('\n', '\\n').replace('\r', '\\r'),
                        clean
                    )
                    data_json = json.loads(repaired)
                segment_list = data_json.get("chunks", [])
                if not segment_list and isinstance(data_json, list):
                    segment_list = data_json
                for c in segment_list:
                    content = c.get("content", "").replace('\\n', '\n')
                    new_chunks.append({
                        "id":               str(uuid.uuid4()),
                        "content":          content,
                        "cliffhanger_note": c.get("cliffhanger_note", ""),
                        "status":           "pending",
                        "word_count":       len(content.split()),
                        "created_at":       now,
                    })
                current_part_count += len(segment_list)
            except json.JSONDecodeError as e:
                print(f"[Serializer] Window {idx+1} JSON error: {e}")
                print(f"[Serializer] Raw (first 400 chars): {response_text[:400]}")
                errors.append(f"Window {idx+1}: {str(e)}")
            except Exception as e:
                print(f"[Serializer] Window {idx+1} error: {e}")
                errors.append(f"Window {idx+1}: {str(e)}")

    if not new_chunks:
        err_detail = "; ".join(errors) if errors else "No chunks were generated."
        return jsonify({"error": f"Serialization failed: {err_detail}"}), 502

    if profile_id:
        work["profile_id"] = profile_id
    work["chunks"].extend(new_chunks)
    work["updated_at"] = now
    write_json(PROMO_WORKS_FILE, works_data)

    for nc in new_chunks:
        create_asset({
            "id":         f"asset_mod_{nc['id']}",
            "type":       "content",
            "title":      f"Segment: {nc['content'][:30]}...",
            "work_id":    work_id,
            "module_id":  nc['id'],
            "production": "done",
        })

    windows_count = 1 if num_chunks else len(windows)
    return jsonify({"chunks": new_chunks, "windows": windows_count, "errors": errors})


@bp.route('/api/works/<work_id>/modules', methods=['POST'])
def create_work_module(work_id):
    data       = request.get_json() or {}
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data["works"] if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    content = data.get('content', '').strip()
    if not content:
        return jsonify({"error": "content is required"}), 400

    new_chunk = {
        "id":               str(uuid.uuid4()),
        "content":          content,
        "cliffhanger_note": data.get('cliffhanger_note', ''),
        "status":           "pending",
        "word_count":       len(content.split()),
        "created_at":       datetime.utcnow().isoformat() + "Z",
    }
    work.setdefault("chunks", []).append(new_chunk)
    write_json(PROMO_WORKS_FILE, works_data)

    create_asset({
        "id":         f"asset_mod_{new_chunk['id']}",
        "type":       "content",
        "title":      f"Manual Segment: {content[:30]}...",
        "work_id":    work_id,
        "module_id":  new_chunk['id'],
        "production": "done",
    })
    return jsonify(new_chunk), 201


@bp.route('/api/works/<work_id>/modules/<module_id>', methods=['PUT', 'DELETE'])
def module_operations(work_id, module_id):
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data["works"] if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    chunk_idx = next((i for i, c in enumerate(work.get("chunks", [])) if c['id'] == module_id), None)
    if chunk_idx is None:
        return jsonify({"error": "Module not found"}), 404

    if request.method == 'DELETE':
        work["chunks"].pop(chunk_idx)
        write_json(PROMO_WORKS_FILE, works_data)
        return jsonify({"ok": True})

    data  = request.get_json() or {}
    chunk = work["chunks"][chunk_idx]
    if 'title' in data:
        chunk['title'] = data['title']
    if 'content' in data:
        chunk['content']    = data['content']
        chunk['word_count'] = len(data['content'].split())
    if 'cliffhanger_note' in data:
        chunk['cliffhanger_note'] = data['cliffhanger_note']
    write_json(PROMO_WORKS_FILE, works_data)
    return jsonify(chunk)


@bp.route('/api/works/<work_id>/modules/<module_id>/queue', methods=['POST'])
def queue_work_module(work_id, module_id):
    from services.scheduler import auto_schedule
    from services.distribution_service import push_to_outbox
    from utils.json_store import read_json as _rj, BASE_DIR as _BASE_DIR

    data       = request.get_json() or {}
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    chunk = next((c for c in work.get("chunks", []) if c['id'] == module_id), None)
    if not chunk:
        return jsonify({"error": "Module not found"}), 404

    if chunk.get("status") == "queued":
        return jsonify({"error": "This module is already in the send queue."}), 409

    # Load settings once
    _ps       = read_json(PROMO_SETTINGS_FILE) or {}
    _wa       = _ps.get('publishing_wa_recipients', {})
    _schedule = _ps.get('delivery_schedule', {})
    _profiles = _ps.get('serializer_profiles', [])

    # Determine if this is a novel work (profile where num_chunks is None = auto/sliding-window)
    work_profile = next((p for p in _profiles if p['id'] == work.get('profile_id', '')), None)
    is_novel = work_profile is not None and work_profile.get('num_chunks') is None

    # VIP group config
    vip_group_id    = _wa.get('vip_group_id', '').strip()
    vip_group_label = _wa.get('vip_group_label', 'VIP WhatsApp Group')

    # Channel config
    _channel_phone = (data.get('recipient_phone') or '').strip() or \
                     _wa.get('channel_id', '').strip() or \
                     os.environ.get('GOWA_CHANNEL_ID', '')
    _channel_name  = (data.get('recipient_name') or '').strip() or \
                     _wa.get('channel_label', 'WA Channel')

    # Determine scheduled_at for the first delivery (VIP or Channel)
    scheduled_at = data.get('scheduled_at')
    if not scheduled_at:
        existing = (_rj(PROMO_MESSAGES_FILE) or {"messages": []}).get("messages", [])
        queued   = [m for m in existing if m.get('status') == 'queued']
        scheduled_at = auto_schedule('story', queued, _schedule)

    # Build message content
    content = chunk["content"]
    post_header = work.get('post_header', '').strip()
    if post_header:
        content = f"{post_header}\n\n{content}"

    cta_url = data.get('cta_url', '').strip()
    if cta_url:
        if cta_url.startswith('/'):
            cta_url = f"https://realmsandroads.com{cta_url}"
        content = content + f"\n\n📖 Read the rest of the story here:\n{cta_url}"

    # Resolve header image — only if user checked "include header image"
    media_url = None
    if data.get('include_header_image'):
        chunk_img = chunk.get('header_image_path', '').strip()
        if chunk_img:
            # Chunk has its own generated image
            filename  = os.path.basename(chunk_img)
            local_img = os.path.join(_BASE_DIR, 'data', 'generated_images', filename)
            media_url = local_img if os.path.exists(local_img) else chunk_img
        elif work.get('pipeline_module_id'):
            # Fall back to pipeline module's chapter header image
            pipeline = read_json(CONTENT_PIPELINE_FILE) or []
            pm = next((e for e in pipeline if e['id'] == work['pipeline_module_id']), None)
            if pm:
                img = (pm.get('assets') or {}).get('header_image_path', '')
                if img:
                    filename  = os.path.basename(img)
                    local_img = os.path.join(_BASE_DIR, 'data', 'generated_images', filename)
                    media_url = local_img if os.path.exists(local_img) else img

    src_ref = {"work_id": work_id, "module_id": module_id}

    if is_novel and vip_group_id:
        # ── Novel dual-delivery: VIP first, Channel 24 h later ────────────────
        vip_result = push_to_outbox(
            recipient_phone = vip_group_id,
            recipient_name  = vip_group_label,
            content         = content,
            source          = "work_serializer",
            scheduled_at    = scheduled_at,
            source_ref      = src_ref,
            media_url       = media_url,
        )

        # Channel slot = VIP + 24 hours
        vip_dt               = datetime.fromisoformat(scheduled_at.replace('Z', '+00:00'))
        channel_scheduled_at = (vip_dt + timedelta(hours=24)).replace(microsecond=0).isoformat()

        channel_result = push_to_outbox(
            recipient_phone = _channel_phone,
            recipient_name  = _channel_name,
            content         = content,
            source          = "work_serializer",
            scheduled_at    = channel_scheduled_at,
            source_ref      = src_ref,
            media_url       = media_url,
        )

        chunk["status"]              = "queued"
        chunk["vip_message_id"]      = vip_result["id"]
        chunk["vip_scheduled_at"]    = scheduled_at
        chunk["channel_message_id"]  = channel_result["id"]
        chunk["channel_scheduled_at"] = channel_scheduled_at
        # Clear any legacy single message_id
        chunk.pop("message_id", None)
        write_json(PROMO_WORKS_FILE, works_data)

        return jsonify({
            "status":               "queued",
            "vip_message_id":       vip_result["id"],
            "channel_message_id":   channel_result["id"],
            "vip_scheduled_at":     scheduled_at,
            "channel_scheduled_at": channel_scheduled_at,
            "ec2_synced":           vip_result["ec2_synced"] and channel_result["ec2_synced"],
        })

    else:
        # ── Non-novel (flash/short story): single delivery to channel ─────────
        result = push_to_outbox(
            recipient_phone = _channel_phone,
            recipient_name  = _channel_name,
            content         = content,
            source          = "work_serializer",
            scheduled_at    = scheduled_at,
            source_ref      = src_ref,
            media_url       = media_url,
        )
        chunk["status"]     = "queued"
        chunk["message_id"] = result["id"]
        write_json(PROMO_WORKS_FILE, works_data)

        return jsonify({
            "id":           result["id"],
            "scheduled_at": scheduled_at,
            "status":       "queued",
            "ec2_synced":   result["ec2_synced"],
        })


@bp.route('/api/works/<work_id>/modules/<module_id>/unqueue', methods=['POST'])
def unqueue_work_module(work_id, module_id):
    """Reset a queued chunk back to pending — does NOT delete messages from EC2."""
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    chunk_idx = next((i for i, c in enumerate(work.get("chunks", [])) if c['id'] == module_id), None)
    if chunk_idx is None:
        return jsonify({"error": "Chunk not found"}), 404

    chunk = work["chunks"][chunk_idx]
    chunk['status'] = 'pending'
    for field in ('message_id', 'vip_message_id', 'channel_message_id',
                  'vip_scheduled_at', 'channel_scheduled_at'):
        chunk.pop(field, None)

    write_json(PROMO_WORKS_FILE, works_data)
    return jsonify({"ok": True})


@bp.route('/api/works/<work_id>/modules/<module_id>/generate-image', methods=['POST'])
def generate_chunk_image(work_id, module_id):
    """Generate a header image for a serializer chunk using AI prompt + Imagen 3."""
    import base64, requests as _req
    from utils.json_store import BASE_DIR as _BASE_DIR

    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    chunk_idx = next((i for i, c in enumerate(work.get("chunks", [])) if c['id'] == module_id), None)
    if chunk_idx is None:
        return jsonify({"error": "Chunk not found"}), 404

    chunk   = work["chunks"][chunk_idx]
    content = chunk.get('content', '').strip()
    if not content:
        return jsonify({"error": "Chunk has no content"}), 400

    sa_path = os.environ.get('GOOGLE_SA_KEY', '')
    if not sa_path or not os.path.exists(sa_path):
        return jsonify({"error": "GOOGLE_SA_KEY not set — required for image generation"}), 503

    try:
        # Step 1: Generate image prompt from chunk content via AI
        image_prompt = call_ai("work_serializer", [
            {"role": "system", "content": HEADER_IMAGE_PROMPT_SYSTEM},
            {"role": "user",   "content": f"Work: {work.get('title', '')}\n\nStory segment:\n{content[:800]}"},
        ], max_tokens=150).strip()

        # Step 2: Call Imagen 3
        from google.oauth2 import service_account as _sa
        from google.auth.transport.requests import Request as _GReq

        creds = _sa.Credentials.from_service_account_file(
            sa_path, scopes=["https://www.googleapis.com/auth/cloud-platform"]
        )
        creds.refresh(_GReq())

        img_resp = _req.post(
            "https://us-central1-aiplatform.googleapis.com/v1/projects/"
            "gen-lang-client-0717388888/locations/us-central1/publishers/google/"
            "models/imagen-3.0-generate-001:predict",
            headers={"Authorization": f"Bearer {creds.token}", "Content-Type": "application/json"},
            json={
                "instances":  [{"prompt": image_prompt + ", no text overlay, cinematic composition"}],
                "parameters": {"sampleCount": 1, "aspectRatio": "16:9", "personGeneration": "allow_all"},
            },
            timeout=90,
        )

        if img_resp.status_code == 429:
            return jsonify({"error": "Imagen rate limit — wait a few seconds and try again."}), 429
        if img_resp.status_code != 200:
            return jsonify({"error": f"Imagen API error {img_resp.status_code}: {img_resp.text[:200]}"}), 502

        predictions = img_resp.json().get("predictions") or []
        if not predictions:
            return jsonify({"error": "Imagen returned no image (safety filter). Try different content."}), 422

        img_bytes = base64.b64decode(predictions[0]["bytesBase64Encoded"])

        # Step 3: Save to disk
        images_dir = os.path.join(_BASE_DIR, 'data', 'generated_images')
        os.makedirs(images_dir, exist_ok=True)
        filename = f"chunk_{module_id}.jpg"
        img_path = os.path.join(images_dir, filename)
        tmp_path = img_path + '.tmp'
        with open(tmp_path, 'wb') as f:
            f.write(img_bytes)
        os.replace(tmp_path, img_path)

        # Step 4: Store on chunk
        image_url = f"/data/images/{filename}"
        chunk['header_image_path'] = image_url
        chunk['image_prompt']      = image_prompt
        write_json(PROMO_WORKS_FILE, works_data)

        return jsonify({"ok": True, "image_url": image_url, "image_prompt": image_prompt})

    except Exception as e:
        return jsonify({"error": f"Image generation failed: {e}"}), 500


# ── Bulk import & asset generation ───────────────────────────────────────────

def _parse_upload_text(request_obj):
    """Extract plain text from multipart file upload or JSON body. Returns (text, error_response)."""
    if request_obj.content_type and 'multipart' in request_obj.content_type:
        uploaded = request_obj.files.get('file')
        if not uploaded:
            return None, (jsonify({"error": "No file provided"}), 400)
        filename = uploaded.filename.lower()
        if filename.endswith('.txt'):
            return uploaded.read().decode('utf-8', errors='replace'), None
        elif filename.endswith('.rtf'):
            raw  = uploaded.read().decode('utf-8', errors='replace')
            text = re.sub(r'\\(par|line)\b\s*', '\n', raw)
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'[{}\\]', '', text)
            return re.sub(r'[ \t]+', ' ', text).strip(), None
        elif filename.endswith('.docx'):
            try:
                from docx import Document as DocxDocument
                uploaded.seek(0)
                doc  = DocxDocument(io.BytesIO(uploaded.read()))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), None
            except Exception as e:
                return None, (jsonify({"error": f"Failed to read DOCX: {e}"}), 400)
        else:
            return None, (jsonify({"error": "Unsupported file type. Use .txt, .rtf, or .docx"}), 400)
    else:
        body = request_obj.get_json() or {}
        return body.get('text', '').strip(), None


@bp.route('/api/works/<work_id>/preview_split', methods=['POST'])
def preview_split(work_id):
    """
    Detect chapter markers in pasted/uploaded text and return a split preview.
    Stores full content in _PENDING_SPLITS[work_id] for the confirm step.
    """
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    text, err = _parse_upload_text(request)
    if err:
        return err
    if not text:
        return jsonify({"error": "No text provided"}), 400

    lines    = text.split('\n')
    chapters = []
    current_title = None
    current_lines = []

    def _flush(title, lines_buf):
        content = '\n'.join(lines_buf).strip()
        if not content:
            return
        idx = len(chapters)
        chapters.append({
            "index":      idx,
            "title":      title or f"Chapter {idx + 1}",
            "content":    content,
            "word_count": len(content.split()),
            "preview":    content[:200],
        })

    for line in lines:
        marker = _detect_chapter_line(line)
        if marker:
            _flush(current_title, current_lines)
            current_title = marker
            current_lines = []
        else:
            current_lines.append(line)

    _flush(current_title, current_lines)

    if not chapters:
        return jsonify({"error": "No content detected after splitting."}), 400

    _PENDING_SPLITS[work_id] = chapters

    # Return preview (no full content — client uses indices to confirm)
    return jsonify({
        "work_id":       work_id,
        "chapter_count": len(chapters),
        "total_words":   sum(c['word_count'] for c in chapters),
        "chapters":      [
            {k: v for k, v in c.items() if k != 'content'}
            for c in chapters
        ],
    })


@bp.route('/api/works/<work_id>/bulk_import_modules', methods=['POST'])
def bulk_import_modules(work_id):
    """
    Confirm a pending split and save chapters as work modules.

    Body: {
      "chapters": [
        {"title": "...", "indices": [0]},          // single chapter
        {"title": "...", "indices": [1, 2]},        // merged chapters
      ]
    }
    Content for each entry comes from _PENDING_SPLITS[work_id] (merged in order).
    """
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    pending = _PENDING_SPLITS.get(work_id)
    if not pending:
        return jsonify({"error": "No pending split found. Call preview_split first."}), 409

    data     = request.get_json() or {}
    chapters = data.get('chapters', [])
    if not chapters:
        return jsonify({"error": "chapters list is required"}), 400

    now        = datetime.utcnow().isoformat() + "Z"
    new_chunks = []

    for ch in chapters:
        indices = ch.get('indices', [])
        title   = (ch.get('title') or '').strip()
        if not indices:
            continue
        # Merge content from the indicated pending split entries
        content_parts = []
        for idx in indices:
            if 0 <= idx < len(pending):
                content_parts.append(pending[idx]['content'])
        content = '\n\n'.join(content_parts).strip()
        if not content:
            continue
        if not title:
            title = pending[indices[0]]['title']

        new_chunk = {
            "id":               str(uuid.uuid4()),
            "title":            title,
            "content":          content,
            "cliffhanger_note": "",
            "status":           "pending",
            "word_count":       len(content.split()),
            "created_at":       now,
        }
        new_chunks.append(new_chunk)

    if not new_chunks:
        return jsonify({"error": "No valid chapters to import"}), 400

    work.setdefault("chunks", []).extend(new_chunks)
    work["updated_at"] = now
    write_json(PROMO_WORKS_FILE, works_data)

    for nc in new_chunks:
        create_asset({
            "id":         f"asset_mod_{nc['id']}",
            "type":       "content",
            "title":      f"{nc['title']}: Content",
            "work_id":    work_id,
            "module_id":  nc['id'],
            "content":    nc['content'],
            "production": "done",
        })

    # Clear pending split
    _PENDING_SPLITS.pop(work_id, None)

    return jsonify({
        "imported": len(new_chunks),
        "chunks":   [{"id": c["id"], "title": c["title"], "word_count": c["word_count"]} for c in new_chunks],
    }), 201


@bp.route('/api/works/<work_id>/bulk_delete_modules', methods=['POST'])
def bulk_delete_modules(work_id):
    """
    Delete all modules (chunks) for a work, along with their assets.
    Blocked if any module has status 'queued' (already in send queue).
    """
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    chunks = work.get("chunks", [])
    queued = [c for c in chunks if c.get("status") == "queued"]
    if queued:
        return jsonify({
            "error": f"{len(queued)} module(s) are currently queued for sending. "
                     "Dequeue or send them before replacing modules.",
            "queued_ids": [c["id"] for c in queued],
        }), 409

    chunk_ids = [c['id'] for c in chunks]

    # Remove assets for these chunks
    if chunk_ids:
        assets_data = read_json(PROMO_ASSETS_FILE) or {"assets": []}
        assets_data["assets"] = [
            a for a in assets_data["assets"]
            if a.get("module_id") not in chunk_ids
        ]
        write_json(PROMO_ASSETS_FILE, assets_data)

    work["chunks"]     = []
    work["updated_at"] = datetime.utcnow().isoformat() + "Z"
    write_json(PROMO_WORKS_FILE, works_data)

    return jsonify({"ok": True, "deleted": len(chunk_ids)})


@bp.route('/api/works/<work_id>/modules/bulk_generate_assets', methods=['POST'])
def bulk_generate_assets(work_id):
    """
    Start a background job to generate synopsis, tagline, and/or blurb for modules.

    Body: {
      "asset_types": ["synopsis", "tagline", "blurb"],  // subset or all
      "module_ids":  ["..."]                             // optional; defaults to all
    }
    Returns: {"job_id": "..."}
    Poll: GET /api/works/jobs/<job_id>
    """
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    from services.asset_manager import list_modules as _list_mods

    data        = request.get_json() or {}
    asset_types = [t for t in data.get('asset_types', ['synopsis', 'tagline', 'blurb'])
                   if t in _ASSET_PROMPTS]
    if not asset_types:
        return jsonify({"error": "No valid asset_types provided. Choose from: synopsis, tagline, blurb"}), 400

    source      = data.get('source', 'chunks')   # 'chunks' (works.json) or 'modules' (modules.json)
    requested   = data.get('module_ids')

    if source == 'modules':
        all_mods     = [m for m in _list_mods() if m.get('work_id') == work_id]
        allowed_ids  = {m['id'] for m in all_mods}
        target_ids   = set(requested) & allowed_ids if requested else allowed_ids
        module_items = [
            {"id": m['id'], "title": m.get('title', ''), "content": m.get('prose', '')}
            for m in all_mods if m['id'] in target_ids
        ]
    else:
        all_chunks   = work.get("chunks", [])
        allowed_ids  = {c['id'] for c in all_chunks}
        target_ids   = set(requested) & allowed_ids if requested else allowed_ids
        module_items = [
            {"id": c['id'],
             "title": c.get('title') or f"Module {i+1}",
             "content": c.get('content', '')}
            for i, c in enumerate(all_chunks) if c['id'] in target_ids
        ]

    if not module_items:
        return jsonify({"error": "No matching modules found"}), 400

    job_id = str(uuid.uuid4())
    BULK_JOBS[job_id] = {
        "job_id":       job_id,
        "work_id":      work_id,
        "status":       "running",
        "asset_types":  asset_types,
        "total":        len(module_items) * len(asset_types),
        "done":         0,
        "current":      "Starting...",
        "errors":       [],
        "completed_at": None,
    }

    t = threading.Thread(
        target=_run_bulk_generate,
        args=(job_id, work_id, module_items, asset_types),
        daemon=True,
    )
    t.start()

    return jsonify({"job_id": job_id, "total": BULK_JOBS[job_id]['total']}), 202


@bp.route('/api/works/jobs/<job_id>', methods=['GET'])
def get_bulk_job(job_id):
    """Poll the status of a bulk generation job."""
    job = BULK_JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)


@bp.route('/api/works/<work_id>/scrivenings', methods=['GET'])
def get_scrivenings(work_id):
    """Return all modules for a work, ordered, with prose from the Content asset."""
    from services.asset_manager import list_modules as _list_mods
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    modules = sorted(
        [m for m in _list_mods() if m.get('work_id') == work_id],
        key=lambda m: m.get('ordinal', 0)
    )

    all_assets = (read_json(PROMO_ASSETS_FILE) or {"assets": []}).get("assets", [])

    enriched = []
    for mod in modules:
        mid = mod['id']
        # Content asset: prefer canonical ID, then any content-type asset for this module
        content_asset = (
            next((a for a in all_assets if a['id'] == f"asset_mod_{mid}"), None)
            or next((a for a in all_assets
                     if a.get('module_id') == mid and a.get('type') == 'content'), None)
        )
        content_asset_id = (content_asset or {}).get('id') or f"asset_mod_{mid}"
        prose = (content_asset or {}).get('content') or mod.get('prose', '')

        enriched.append({
            **mod,
            'prose':            prose,
            'content_asset_id': content_asset_id,
        })

    return jsonify({
        "work_id":    work_id,
        "work_title": work.get('title', ''),
        "modules":    enriched,
    })


@bp.route('/api/works/<work_id>/bulk_import_chapters', methods=['POST'])
def bulk_import_chapters(work_id):
    """
    Import confirmed chapters into modules.json (used by the Inventory tab).
    Shares _PENDING_SPLITS with preview_split.

    Body same as bulk_import_modules: {chapters: [{title, indices}]}
    """
    from services.asset_manager import save_module as _save_mod

    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    pending = _PENDING_SPLITS.get(work_id)
    if not pending:
        return jsonify({"error": "No pending split found. Call preview_split first."}), 409

    data     = request.get_json() or {}
    chapters = data.get('chapters', [])
    if not chapters:
        return jsonify({"error": "chapters list is required"}), 400

    now         = datetime.utcnow().isoformat() + "Z"
    new_modules = []

    for ordinal, ch in enumerate(chapters):
        indices = ch.get('indices', [])
        title   = (ch.get('title') or '').strip()
        if not indices:
            continue
        content_parts = [pending[i]['content'] for i in indices if 0 <= i < len(pending)]
        prose = '\n\n'.join(content_parts).strip()
        if not prose:
            continue
        if not title:
            title = pending[indices[0]]['title']

        module_id = str(uuid.uuid4())
        module = {
            "id":         module_id,
            "title":      title,
            "prose":      prose,
            "status":     "draft",
            "work_id":    work_id,
            "ordinal":    ordinal,
            "created_at": now,
            "updated_at": now,
        }
        _save_mod(module)
        create_asset({
            "id":         f"asset_mod_{module_id}",
            "type":       "content",
            "title":      f"{title}: Content",
            "work_id":    work_id,
            "module_id":  module_id,
            "content":    prose,
            "production": "done",
        })
        new_modules.append({"id": module_id, "title": title,
                            "word_count": len(prose.split())})

    if not new_modules:
        return jsonify({"error": "No valid chapters to import"}), 400

    _PENDING_SPLITS.pop(work_id, None)
    return jsonify({"imported": len(new_modules), "modules": new_modules}), 201


@bp.route('/api/works/<work_id>/bulk_delete_chapters', methods=['POST'])
def bulk_delete_chapters(work_id):
    """
    Delete all modules (modules.json) for a work + their assets.
    Blocked if any queued outbox messages reference this work.
    """
    from services.asset_manager import list_modules as _list_mods, delete_module as _del_mod

    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    modules = [m for m in _list_mods() if m.get('work_id') == work_id]

    msgs_data = read_json(PROMO_MESSAGES_FILE) or {"messages": []}
    queued = [
        m for m in msgs_data["messages"]
        if m.get('status') in ('queued', 'dispatched')
        and m.get('source_ref', {}).get('work_id') == work_id
    ]
    if queued:
        return jsonify({
            "error": f"{len(queued)} outbox message(s) are queued for this work. "
                     "Send or cancel them before replacing chapters.",
        }), 409

    deleted = 0
    for mod in modules:
        _del_mod(mod['id'])
        deleted += 1

    return jsonify({"ok": True, "deleted": deleted})


@bp.route('/api/works/<work_id>/modules/bulk_generate_header_prompts', methods=['POST'])
def bulk_generate_header_prompts(work_id):
    """
    Generate header_image_prompt assets for modules that already have synopses.
    Optionally accepts a reference image (base64) to guide visual style.

    Body: {
      "module_ids":        [...],          // optional; defaults to all with synopses
      "reference_image_b64": "...",        // optional base64 JPEG/PNG
      "reference_image_mime": "image/jpeg" // default image/jpeg
    }
    Uses Gemini Flash (Vertex AI) if reference_image_b64 provided; else DeepSeek/OpenAI.
    Runs synchronously (typically ≤20 modules — manageable).
    """
    works_data = read_json(PROMO_WORKS_FILE) or {"works": []}
    work = next((w for w in works_data.get("works", []) if w['id'] == work_id), None)
    if not work:
        return jsonify({"error": "Work not found"}), 404

    data           = request.get_json() or {}
    ref_image_b64  = data.get('reference_image_b64', '')
    ref_image_mime = data.get('reference_image_mime', 'image/jpeg')
    requested_ids  = data.get('module_ids')

    all_chunks = work.get("chunks", [])
    if requested_ids:
        all_chunks = [c for c in all_chunks if c['id'] in requested_ids]

    # Gate on synopsis existing
    eligible = []
    for chunk in all_chunks:
        synopsis_asset = get_asset(f"asset_synopsis_{chunk['id']}")
        if synopsis_asset and synopsis_asset.get('content', '').strip():
            eligible.append((chunk, synopsis_asset['content']))

    if not eligible:
        return jsonify({
            "error": "No modules with synopses found. Generate synopses first."
        }), 409

    # If reference image provided, use Gemini Flash to describe its style
    style_guide = ''
    if ref_image_b64:
        sa_path = os.environ.get('GOOGLE_SA_KEY', '')
        if not sa_path or not os.path.exists(sa_path):
            return jsonify({"error": "GOOGLE_SA_KEY not set — required for reference image processing"}), 503
        try:
            import requests as _req
            from google.oauth2 import service_account as _sa
            from google.auth.transport.requests import Request as _Req

            _creds = _sa.Credentials.from_service_account_file(
                sa_path, scopes=["https://www.googleapis.com/auth/cloud-platform"]
            )
            _creds.refresh(_Req())

            gemini_url = (
                "https://us-central1-aiplatform.googleapis.com/v1/projects/"
                "gen-lang-client-0717388888/locations/us-central1/publishers/google/"
                "models/gemini-2.0-flash-001:generateContent"
            )
            vision_body = {
                "contents": [{
                    "role": "user",
                    "parts": [
                        {"inlineData": {"mimeType": ref_image_mime, "data": ref_image_b64}},
                        {"text": (
                            "Describe the visual style and aesthetic of this image in 40-60 words. "
                            "Focus on: lighting quality, colour palette, mood, photographic style, "
                            "composition. This description will be used as a style guide for AI image generation."
                        )},
                    ],
                }],
            }
            resp = _req.post(
                gemini_url,
                headers={"Authorization": f"Bearer {_creds.token}", "Content-Type": "application/json"},
                json=vision_body,
                timeout=30,
            )
            if resp.status_code == 200:
                candidates = resp.json().get("candidates", [])
                if candidates:
                    style_guide = candidates[0]["content"]["parts"][0].get("text", "").strip()
        except Exception as e:
            # Non-fatal: proceed without style guide
            style_guide = ''
            print(f"[HeaderPrompt] Gemini vision failed: {e}")

    generated = 0
    skipped   = 0
    errors    = []

    for chunk, synopsis in eligible:
        asset_id = f"asset_header_image_prompt_{chunk['id']}"
        existing = get_asset(asset_id)
        if existing and existing.get('content', '').strip():
            skipped += 1
            continue

        title = chunk.get('title') or f"Module {chunk['id'][:8]}"
        user_parts = [f"Chapter Title: {title}", f"Synopsis: {synopsis}"]
        if style_guide:
            user_parts.append(f"Visual Style Reference: {style_guide}")
        user_msg = "\n\n".join(user_parts)

        try:
            result = call_ai("work_serializer", [
                {"role": "system", "content": HEADER_IMAGE_PROMPT_SYSTEM},
                {"role": "user",   "content": user_msg},
            ], max_tokens=200)

            create_asset({
                "id":         asset_id,
                "type":       "header_image_prompt",
                "title":      f"{title} \u2014 Header Image Prompt",
                "work_id":    work_id,
                "module_id":  chunk['id'],
                "content":    result.strip(),
                "production": "done",
            })
            generated += 1
        except Exception as e:
            errors.append(f"{title[:30]}: {str(e)}")

    return jsonify({
        "generated": generated,
        "skipped":   skipped,
        "errors":    errors,
    })


# ── Backwards-compat aliases for old /api/promo/books/* endpoints ────────────

@bp.route('/api/promo/books', methods=['GET'])
def list_works_compat():
    data = read_json(PROMO_WORKS_FILE) or {"works": []}
    _sync_chunk_statuses(data)
    return jsonify({"books": data.get("works", [])})

@bp.route('/api/promo/books', methods=['POST'])
def create_work_compat():
    return create_work()

@bp.route('/api/promo/books/<work_id>', methods=['GET'])
def get_work_compat(work_id):
    return get_work(work_id)

@bp.route('/api/promo/books/<work_id>', methods=['PUT'])
def update_work_compat(work_id):
    return update_work(work_id)

@bp.route('/api/promo/books/<work_id>', methods=['DELETE'])
def delete_work_compat(work_id):
    return delete_work_route(work_id)

@bp.route('/api/promo/books/<work_id>/ingest', methods=['POST'])
def ingest_compat(work_id):
    return ingest_work_content(work_id)

@bp.route('/api/promo/books/<work_id>/chunks', methods=['POST'])
def create_chunk_compat(work_id):
    return create_work_module(work_id)

@bp.route('/api/promo/books/<work_id>/chunks/<module_id>', methods=['PUT', 'DELETE'])
def chunk_ops_compat(work_id, module_id):
    return module_operations(work_id, module_id)

@bp.route('/api/promo/books/<work_id>/chunks/<module_id>/queue', methods=['POST'])
def queue_chunk_compat(work_id, module_id):
    return queue_work_module(work_id, module_id)
